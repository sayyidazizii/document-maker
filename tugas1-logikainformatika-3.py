from docx import Document

# Membuat dokumen baru
doc = Document()
doc.add_heading('Analisis Kalimat Proposisional', level=1)

# Menambahkan deskripsi kalimat
doc.add_paragraph(
    "Diberikan kalimat proposisional:\n"
    "F: if (not R) then (not P if and only if (R and P))\n\n"
    "Kalimat ini menyatakan suatu implikasi, di mana bagian 'if' adalah premis dan bagian 'then' adalah konsekuensi.\n"
)

# Menjelaskan sifat kalimat
doc.add_heading('Sifat Kalimat', level=2)
doc.add_paragraph(
    "Kalimat proposisional F dapat dikatakan sebagai sebuah implikasi yang menyatakan bahwa jika R tidak benar (¬R), "
    "maka ¬P ekuivalen dengan (R ∧ P). Untuk menunjukkan sifat kalimat ini, kita akan menggunakan tabel kebenaran.\n"
)

# Menambahkan tabel kebenaran
doc.add_heading('Tabel Kebenaran untuk F', level=2)
table = doc.add_table(rows=1, cols=11)
hdr_cells = table.rows[0].cells
headers = ['P', 'R', '¬R', 'R ∧ P', '¬P', '¬P ↔ (R ∧ P)', 'F: ¬R → (¬P ↔ (R ∧ P))']
for i in range(len(headers)):
    hdr_cells[i].text = headers[i]

# Data tabel kebenaran
data = [
    ('T', 'T', 'F', 'T', 'F', 'T', 'T'),
    ('T', 'F', 'T', 'F', 'F', 'F', 'F'),
    ('F', 'T', 'F', 'F', 'T', 'F', 'T'),
    ('F', 'F', 'T', 'F', 'T', 'F', 'F'),
]

# Menambahkan data ke dalam tabel
for row in data:
    cells = table.add_row().cells
    for i in range(len(row)):
        cells[i].text = row[i]

# Menambahkan penjelasan pohon semantik
doc.add_heading('Pohon Semantik', level=2)
doc.add_paragraph(
    "Pohon semantik untuk kalimat F dapat digambarkan sebagai berikut:\n"
)
doc.add_paragraph(
    "          ¬R\n"
    "          / \\\n"
    "         T   F\n"
    "        /     \\\n"
    "    ¬P ↔ (R ∧ P)\n"
    "      /   |    \\\n"
    "     T    F    T\n"
    "   (F)   (T)   (F)\n"
)

# Menyimpulkan analisis
doc.add_heading('Kesimpulan', level=2)
doc.add_paragraph(
    "Dari tabel kebenaran dan pohon semantik, kita dapat melihat bahwa:\n"
    "- Kalimat F tidak selalu benar; ia hanya benar pada beberapa kombinasi nilai kebenaran.\n"
    "- Sifat: F adalah implikasi yang tidak selalu benar. Dalam konteks logika, kita dapat menyatakan bahwa kalimat ini "
    "dapat digunakan untuk menunjukkan hubungan antara R dan P di mana satu variabel mempengaruhi yang lain melalui "
    "negasi dan konjungsi."
)

# Menyimpan dokumen ke file
doc.save('D:/var/www/html/python/document-maker/tugas1-logikainformatika-3.docx')

print("Dokumen berhasil disimpan.")
