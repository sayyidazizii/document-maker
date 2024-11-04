from docx import Document

# Membuat dokumen baru
doc = Document()

# Menambahkan judul
doc.add_heading('Analisis Kasus Lazada', level=1)

# Menambahkan konten
content = [
    ("1. Alasan dan Dampak Penerapan CSR bagi Perusahaan", [
        "Alasan:",
        "- Membangun Citra Positif: CSR membantu perusahaan membangun citra positif di mata masyarakat.",
        "- Mengurangi Risiko: Melalui program CSR, perusahaan dapat mengurangi potensi konflik dengan masyarakat.",
        "- Dukungan dari Masyarakat: Dengan melibatkan masyarakat dalam program CSR, perusahaan dapat memperoleh dukungan dari warga sekitar.",
        "Dampak:",
        "- Peningkatan Loyalitas Pelanggan: Ketika masyarakat melihat perusahaan berkomitmen, mereka lebih cenderung mendukung.",
        "- Dampak Sosial Positif: Penerapan CSR dapat menciptakan dampak sosial yang positif.",
        "- Peningkatan Kinerja Keuangan: Banyak penelitian menunjukkan bahwa perusahaan yang menerapkan CSR dapat mencapai kinerja keuangan yang lebih baik."
    ]),
    ("2. Alasan Perusahaan Kontra dalam Menerapkan CSR", [
        "- Biaya Tambahan: CSR dapat dianggap sebagai biaya tambahan yang mengurangi laba jangka pendek.",
        "- Fokus pada Keuntungan: Beberapa perusahaan memiliki fokus utama pada keuntungan.",
        "- Kurangnya Pemahaman: Terkadang, perusahaan tidak memahami manfaat jangka panjang dari CSR."
    ]),
    ("3. Level Penerapan CSR di Lazada", [
        "Penerapan CSR di Lazada tampaknya berada pada level yang rendah.",
        "Hal ini terlihat dari kurangnya respons terhadap proposal yang diajukan oleh masyarakat."
    ]),
    ("4. Perencanaan Strategis jika Saya CEO Lazada", [
        "- Evaluasi Program CSR Saat Ini: Lakukan audit menyeluruh terhadap program CSR.",
        "- Membangun Kemitraan dengan Komunitas: Libatkan masyarakat dalam perencanaan program CSR.",
        "- Menetapkan Anggaran yang Layak untuk CSR: Alokasikan dana yang lebih besar untuk program CSR.",
        "- Transparansi dan Akuntabilitas: Publikasikan laporan tahunan tentang kegiatan CSR.",
        "- Kampanye Pemasaran Sosial: Gunakan program CSR sebagai bagian dari strategi pemasaran.",
        "- Inisiatif Jangka Panjang: Kembangkan program-program jangka panjang yang dapat memberdayakan masyarakat."
    ])
]

# Menambahkan konten ke dokumen
for title, items in content:
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(item)

# Menyimpan dokumen
doc.save('D:/var/www/html/python/document-maker/tugas1-manajemen.docx')
