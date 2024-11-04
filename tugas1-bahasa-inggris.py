from docx import Document

# Buat dokumen baru
doc = Document()

# Menambahkan judul
doc.add_heading('Jawaban Tugas', level=1)

# Menambahkan isi jawaban dalam bahasa Indonesia
doc.add_heading('A. Identifikasi ragam bahasa apakah yang digunakan dalam percakapan tersebut.', level=2)
doc.add_paragraph("Ragam bahasa yang digunakan dalam percakapan tersebut adalah bahasa informal atau bahasa santai.")

doc.add_heading('B. Apa alasan Anda untuk jawaban A?', level=2)
doc.add_paragraph(
    "Alasan untuk jawaban A adalah karena percakapan ini menunjukkan gaya bahasa yang akrab dan tidak formal, di mana kedua orang "
    "yang terlibat (Amar dan Alex) menggunakan istilah sehari-hari dan frasa yang umum digunakan dalam percakapan kasual. "
    "Contohnya, penggunaan sapaan seperti 'Hey' dan pertanyaan yang langsung dan santai seperti 'What’s up?' dan 'Sounds fun!' "
    "menunjukkan suasana yang tidak kaku. Selain itu, penggunaan kata-kata seperti 'chilling,' 'cool,' dan 'just bring yourself!' "
    "menunjukkan keakraban dan kenyamanan dalam berbicara satu sama lain."
)

doc.add_heading('C. Identifikasi topik percakapan tersebut.', level=2)
doc.add_paragraph("Topik percakapan tersebut adalah rencana untuk mengadakan pertemuan santai atau acara kumpul-kumpul di rumah Amar pada akhir pekan.")

# Menambahkan isi jawaban dalam bahasa Inggris
doc.add_heading('A. Identify what style of language used in the conversation.', level=2)
doc.add_paragraph("The style of language used in the conversation is informal language or casual language.")

doc.add_heading('B. What is the reason for your answer in question A?', level=2)
doc.add_paragraph(
    "The reason for answer A is that the conversation shows a friendly and informal style of language, where both people involved "
    "(Amar and Alex) use everyday terms and phrases commonly used in casual conversations. For example, the use of greetings like 'Hey' "
    "and direct, relaxed questions like 'What’s up?' and 'Sounds fun!' indicate a relaxed atmosphere. Additionally, the use of words "
    "like 'chilling,' 'cool,' and 'just bring yourself!' demonstrates familiarity and comfort in speaking with each other."
)

doc.add_heading('C. Identify the topic of the conversation.', level=2)
doc.add_paragraph("The topic of the conversation is the plan to have a casual gathering or get-together at Amar's house over the weekend.")

# Simpan dokumen
doc.save('D:/var/www/html/python/document-maker/tugas1-bahasa-inggris.docx')
