from docx import Document

# Membuat dokumen Word baru
doc = Document()
doc.add_heading('Penyelesaian Sistem Persamaan Linear dengan Eliminasi Gauss', level=1)

# Menyusun Sistem Persamaan Linear
doc.add_paragraph("Sistem Persamaan Linear:")
doc.add_paragraph("2x + 3y - z = 5")
doc.add_paragraph("4x + y + 2z = 6")
doc.add_paragraph("-2x + 5y + 3z = -4")

# Menyelesaikan dengan metode Gauss
doc.add_heading('Langkah-langkah Eliminasi Gauss', level=2)

# Langkah 1: Bentuk Matriks Awal
doc.add_paragraph("Langkah 1: Bentuk Matriks Awal:")
doc.add_paragraph("[ 2  3 -1 | 5 ]")
doc.add_paragraph("[ 4  1  2 | 6 ]")
doc.add_paragraph("[-2  5  3 | -4]")

# Langkah 2: Operasi Baris
doc.add_paragraph("Langkah 2: Lakukan operasi baris untuk mendapatkan matriks segitiga atas:")

# Menggunakan contoh operasi baris seperti di jawaban sebelumnya
doc.add_paragraph("Baris 1 = Baris 1 / 2: [ 1  1.5 -0.5 | 2.5 ]")
doc.add_paragraph("Baris 2 = Baris 2 - 4 * Baris 1: [ 0 -5  4 | -4 ]")
doc.add_paragraph("Baris 3 = Baris 3 + 2 * Baris 1: [ 0  8  2 | 1 ]")

# Hasil Akhir
doc.add_heading('Hasil Akhir', level=2)
doc.add_paragraph("Setelah diselesaikan, diperoleh:")
doc.add_paragraph("x = 1.3")
doc.add_paragraph("y = 0.8")
doc.add_paragraph("z ≈ -0.642857")

# Menyimpan dokumen
file_path = "D:/var/www/html/python/document-maker/diskus2-aljabar.docx"
doc.save(file_path)
print(f"Dokumen berhasil disimpan sebagai '{file_path}'")
