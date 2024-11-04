from docx import Document

# Membuat dokumen Word baru
doc = Document()

# Menambahkan judul
doc.add_heading('Analisis Nilai Kebenaran Kalimat Logika Proposisional', level=1)

# Menambahkan isi dokumen
doc.add_paragraph(
    "Untuk menentukan apakah nilai kebenaran kalimat "
    "F: ((P ∧ Q) ∨ S) ↔ (P ∧ (Q ∨ S)) berubah di bawah interpretasi yang diperluas, "
    "kita perlu mengevaluasi kalimat F di bawah dua interpretasi yang berbeda:\n\n"
    "1. Interpretasi awal: J: {P ← false, Q ← true, S ← true}\n"
    "2. Interpretasi yang diperluas: <S ← false> • <Q ← false> • J, "
    "di mana P = false, Q = false, S = false.\n\n"
    "Mari kita evaluasi F di bawah kedua interpretasi ini."
)

# Menambahkan subjudul dan isi interpretasi awal
doc.add_heading('1. Evaluasi F di bawah Interpretasi Awal J', level=2)
doc.add_paragraph(
    "- Ekspresi (P ∧ Q) ∨ S:\n"
    "  - P = false, Q = true, dan S = true.\n"
    "  - (P ∧ Q) = false ∧ true = false.\n"
    "  - (P ∧ Q) ∨ S = false ∨ true = true.\n\n"
    "- Ekspresi P ∧ (Q ∨ S):\n"
    "  - (Q ∨ S) = true ∨ true = true.\n"
    "  - P ∧ (Q ∨ S) = false ∧ true = false.\n\n"
    "Jadi, F = ((P ∧ Q) ∨ S) ↔ (P ∧ (Q ∨ S)) = true ↔ false = false."
)

# Menambahkan subjudul dan isi interpretasi diperluas
doc.add_heading('2. Evaluasi F di bawah Interpretasi yang Diperluas', level=2)
doc.add_paragraph(
    "- Ekspresi (P ∧ Q) ∨ S:\n"
    "  - P = false, Q = false, dan S = false.\n"
    "  - (P ∧ Q) = false ∧ false = false.\n"
    "  - (P ∧ Q) ∨ S = false ∨ false = false.\n\n"
    "- Ekspresi P ∧ (Q ∨ S):\n"
    "  - (Q ∨ S) = false ∨ false = false.\n"
    "  - P ∧ (Q ∨ S) = false ∧ false = false.\n\n"
    "Jadi, F = ((P ∧ Q) ∨ S) ↔ (P ∧ (Q ∨ S)) = false ↔ false = true."
)

# Menambahkan kesimpulan
doc.add_heading('Kesimpulan', level=2)
doc.add_paragraph(
    "Nilai kebenaran kalimat F berubah di bawah interpretasi yang diperluas. "
    "Di bawah interpretasi awal J, F bernilai false, sedangkan di bawah interpretasi "
    "yang diperluas, F bernilai true."
)

# Menyimpan dokumen
file_path = "D:/var/www/html/python/document-maker/Analisis_Nilai_Kebenaran_Kalimat_Logika_Proposisional.docx"

doc.save(file_path)

file_path
