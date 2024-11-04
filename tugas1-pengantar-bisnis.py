from docx import Document

# Membuat dokumen Word
doc = Document()

# Menambahkan judul
doc.add_heading('Tugas Pengantar Bisnis', level=1)

# Menambahkan pertanyaan dan jawaban
# Pertanyaan 1
doc.add_heading('1. Peringkat Indeks Kemudahan Berbisnis', level=2)
doc.add_paragraph(
    "Indeks Kemudahan Berbisnis (Ease of Doing Business Index) adalah indikator yang digunakan untuk menilai lingkungan bisnis suatu negara berdasarkan "
    "berbagai aspek, seperti regulasi, infrastruktur, dan perlindungan hukum. Indeks ini disusun oleh Bank Dunia dan mencakup berbagai indikator "
    "yang menggambarkan kemudahan melakukan bisnis di negara tertentu.\n\n"
    "Beberapa aspek yang dinilai dalam indeks ini meliputi:\n"
    "- Proses mendirikan usaha: Kemudahan dan kecepatan dalam mengurus izin usaha.\n"
    "- Mendapatkan izin bangunan: Proses pengajuan dan penerimaan izin untuk membangun infrastruktur.\n"
    "- Perlindungan investor: Tingkat perlindungan hukum untuk investor, termasuk transparansi dan akuntabilitas perusahaan.\n"
    "- Pengadaan barang dan jasa: Proses pengadaan barang dan jasa oleh pemerintah.\n"
    "- Perdagangan lintas batas: Kemudahan dalam ekspor dan impor barang.\n\n"
    "Peringkat negara berdasarkan indeks ini sering berubah setiap tahunnya, dan negara-negara yang memiliki peringkat tinggi biasanya memiliki "
    "regulasi yang lebih sederhana, infrastruktur yang baik, dan perlindungan hukum yang kuat. Misalnya, negara seperti Singapura, Selandia Baru, "
    "dan Denmark sering menempati posisi teratas dalam indeks ini."
)

# Pertanyaan 2
doc.add_heading('2. Langkah-langkah Mendirikan Perseroan Terbatas', level=2)
doc.add_paragraph(
    "Untuk mendirikan perseroan terbatas (PT), berikut adalah langkah-langkah yang dapat dilakukan:\n"
    "1. Penyusunan Akta Pendirian: Membuat akta pendirian yang memuat informasi mengenai nama perusahaan, maksud dan tujuan, modal dasar, "
    "serta struktur organisasi. Akta ini harus dibuat di hadapan notaris.\n"
    "2. Pengajuan Permohonan Nama: Mengajukan permohonan untuk mendapatkan persetujuan nama perusahaan kepada Kementerian Hukum dan HAM agar "
    "nama yang diusulkan tidak sama dengan nama perusahaan lain.\n"
    "3. Penetapan Modal Dasar: Menentukan modal dasar perusahaan yang terdiri dari modal disetor dan modal yang belum disetor sesuai ketentuan UU PT.\n"
    "4. Mendapatkan Nomor Induk Berusaha (NIB): Mendaftar di sistem Online Single Submission (OSS) untuk mendapatkan NIB sebagai identitas perusahaan.\n"
    "5. Pendaftaran Ke Kementerian Hukum dan HAM: Mengajukan pendaftaran akta pendirian perusahaan ke Kementerian Hukum dan HAM untuk mendapatkan "
    "pengesahan badan hukum.\n"
    "6. Pengurusan Izin Usaha dan Izin Operasional: Mengurus izin usaha dan izin operasional sesuai dengan jenis usaha yang akan dijalankan, "
    "seperti SIUP (Surat Izin Usaha Perdagangan), TDP (Tanda Daftar Perusahaan), dan lainnya.\n"
    "7. Pembukaan Rekening Bank: Membuka rekening bank atas nama perusahaan untuk keperluan transaksi bisnis.\n"
    "8. Pengurusan NPWP (Nomor Pokok Wajib Pajak): Mendaftarkan perusahaan untuk mendapatkan NPWP sebagai identitas perpajakan."
)

# Pertanyaan 3
doc.add_heading('3. Perkembangan Bisnis pada Era 4.0', level=2)
doc.add_paragraph(
    "Era 4.0, atau Revolusi Industri 4.0, ditandai oleh digitalisasi dan integrasi teknologi informasi dalam berbagai aspek kehidupan, termasuk bisnis. "
    "Berikut adalah beberapa perkembangan bisnis yang terjadi pada era 4.0:\n"
    "1. Digitalisasi dan Automatisasi: Banyak perusahaan beralih ke platform digital untuk meningkatkan efisiensi operasional, seperti penggunaan perangkat "
    "lunak manajemen bisnis, otomatisasi proses, dan sistem berbasis cloud.\n"
    "2. E-commerce: Pertumbuhan pesat dalam perdagangan elektronik memungkinkan bisnis menjangkau pelanggan global dengan lebih mudah, serta memberikan "
    "pengalaman berbelanja yang lebih nyaman melalui platform online.\n"
    "3. Analisis Data Besar (Big Data): Penggunaan data besar untuk menganalisis perilaku konsumen, memprediksi tren pasar, dan mengoptimalkan strategi "
    "pemasaran. Perusahaan dapat membuat keputusan yang lebih baik berdasarkan data yang terkumpul.\n"
    "4. Internet of Things (IoT): Penggunaan perangkat terhubung yang dapat berkomunikasi satu sama lain untuk meningkatkan efisiensi operasional, seperti "
    "dalam manajemen rantai pasokan dan pemantauan aset.\n"
    "5. Kecerdasan Buatan (AI): Penerapan AI dalam berbagai bidang bisnis, seperti customer service melalui chatbot, personalisasi pengalaman pengguna, "
    "dan pengambilan keputusan berbasis machine learning.\n"
    "6. Keberlanjutan dan Tanggung Jawab Sosial: Meningkatnya kesadaran akan isu keberlanjutan mendorong perusahaan untuk menerapkan praktik bisnis yang "
    "ramah lingkungan dan bertanggung jawab sosial."
)

# Menyimpan dokumen
doc.save('D:/var/www/html/python/document-maker/tugas1-pengantar-bisnis.docx')
