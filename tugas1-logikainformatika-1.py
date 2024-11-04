from docx import Document

# Membuat dokumen baru
doc = Document()
doc.add_heading('Hasil Interpretasi Logika Proposisional', level=1)

# Menambahkan deskripsi
doc.add_paragraph(
    "Diketahui dua kalimat logika proposisional:\n"
    "F: if true then (not P and Q) else (Q or not S)\n"
    "G: not (P or Q) and (if P then R else (P and not Q)).\n\n"
    "Kalimat F akan bernilai true ketika:\n"
    "- P adalah false (F) dan Q adalah true (T).\n\n"
    "Kalimat G akan bernilai true ketika:\n"
    "- P dan Q adalah false (F) (nilai R tidak berpengaruh).\n\n"
    "Berikut adalah tabel hasil interpretasi logika proposisional F dan G."
)

# Menambahkan tabel untuk hasil
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'P'
hdr_cells[1].text = 'Q'
hdr_cells[2].text = 'R'
hdr_cells[3].text = 'F (not P and Q)'
hdr_cells[4].text = 'G (not (P or Q) and ...)'

# Data interpretasi
data = [
    ('T', 'T', 'T', 'F', 'F'),
    ('T', 'T', 'F', 'F', 'F'),
    ('T', 'F', 'T', 'F', 'F'),
    ('T', 'F', 'F', 'F', 'F'),
    ('F', 'T', 'T', 'F', 'F'),
    ('F', 'T', 'F', 'F', 'F'),
    ('F', 'F', 'T', 'T', 'T'),
    ('F', 'F', 'F', 'T', 'T'),
]

# Menambahkan data ke dalam tabel
for row in data:
    cells = table.add_row().cells
    for i in range(len(row)):
        cells[i].text = row[i]

# Menyimpan dokumen ke file
doc.save('D:/var/www/html/python/document-maker/tugas1-logikainformatika-1.docx')

print("Dokumen berhasil disimpan.")
