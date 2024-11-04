from docx import Document

# Buat dokumen Word baru
doc = Document()
doc.add_heading('Pohon Semantik untuk Kalimat Logika F', level=1)

# Tambahkan deskripsi
doc.add_paragraph('Berikut adalah pohon semantik untuk kalimat logika:')
doc.add_paragraph('F: if (not R or P) then (P or R) else (not R and Q)')
doc.add_paragraph(' ')
doc.add_paragraph('Pohon semantik dapat digambarkan sebagai berikut:')
doc.add_paragraph('1. F')
doc.add_paragraph('   ├─ if (not R or P)')
doc.add_paragraph('   │      ├─ not R')
doc.add_paragraph('   │      └─ P')
doc.add_paragraph('   └─ else (not R and Q)')
doc.add_paragraph('          ├─ not R')
doc.add_paragraph('          └─ Q')
doc.add_paragraph(' ')

# Tambahkan bagian kesimpulan
doc.add_heading('Kesimpulan', level=2)
doc.add_paragraph(
    "Kalimat logika yang diberikan menghasilkan pohon semantik yang menggambarkan "
    "relasi antar kondisi dan hasil. Berdasarkan analisis, kita dapat menyimpulkan "
    "bahwa kalimat ini adalah kontingen dan bukan tautologi, karena nilai kebenaran "
    "F tergantung pada kombinasi nilai P, Q, dan R."
)

# Simpan dokumen
doc.save('D:/var/www/html/python/document-maker/semantic_tree_f.docx')

