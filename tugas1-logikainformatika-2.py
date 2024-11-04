from docx import Document

# Membuat dokumen baru
doc = Document()
doc.add_heading('Ekuivalensi Kalimat Proposisional', level=1)

# Menambahkan deskripsi
doc.add_paragraph(
    "Diberikan dua kalimat logika proposisional:\n"
    "F: if (P and Q) then R\n"
    "G: if (not R) then (not P or not Q)\n\n"
    "Untuk menunjukkan bahwa kalimat F dan G ekuivalen, kita dapat menggunakan tabel kebenaran.\n"
    "Dua pernyataan proposisional dikatakan ekuivalen jika mereka memiliki nilai kebenaran yang sama untuk semua kemungkinan "
    "kombinasi nilai dari variabel-variabelnya."
)

# Menambahkan tabel kebenaran
table = doc.add_table(rows=1, cols=11)
hdr_cells = table.rows[0].cells
headers = ['P', 'Q', 'R', 'P ∧ Q', 'F: (P ∧ Q) → R', '¬R', '¬P', '¬Q', '¬P ∨ ¬Q', 'G: ¬R → (¬P ∨ ¬Q)', 'Ekuivalen']
for i in range(len(headers)):
    hdr_cells[i].text = headers[i]

# Data tabel kebenaran
data = [
    ('T', 'T', 'T', 'T', 'T', 'F', 'F', 'F', 'F', 'T', 'Ekuivalen'),
    ('T', 'T', 'F', 'T', 'F', 'T', 'F', 'F', 'F', 'F', 'Tidak Ekuivalen'),
    ('T', 'F', 'T', 'F', 'T', 'F', 'F', 'T', 'T', 'T', 'Ekuivalen'),
    ('T', 'F', 'F', 'F', 'T', 'T', 'F', 'T', 'T', 'T', 'Ekuivalen'),
    ('F', 'T', 'T', 'F', 'T', 'F', 'T', 'F', 'T', 'T', 'Ekuivalen'),
    ('F', 'T', 'F', 'F', 'T', 'T', 'T', 'F', 'T', 'T', 'Ekuivalen'),
    ('F', 'F', 'T', 'F', 'T', 'F', 'T', 'T', 'T', 'T', 'Ekuivalen'),
    ('F', 'F', 'F', 'F', 'T', 'T', 'T', 'T', 'T', 'T', 'Ekuivalen'),
]

# Menambahkan data ke dalam tabel
for row in data:
    cells = table.add_row().cells
    for i in range(len(row)):
        cells[i].text = row[i]

# Menambahkan kesimpulan
doc.add_paragraph(
    "\nKesimpulan:\n"
    "Berdasarkan tabel kebenaran di atas, kita dapat melihat bahwa nilai kebenaran dari kalimat F dan G sama pada semua "
    "kombinasi nilai P, Q, dan R. Oleh karena itu, kalimat F dan G dapat dinyatakan sebagai ekuivalen:\n"
    "F ≡ G"
)

# Menyimpan dokumen ke file
doc.save('D:/var/www/html/python/document-maker/tugas1-logikainformatika-2.docx')

print("Dokumen berhasil disimpan.")
